from __future__ import annotations

import json
import math
import re
import shutil
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Any

import fitz
import matplotlib.pyplot as plt
import pandas as pd

try:
    from docx import Document  # type: ignore
except Exception:  # pragma: no cover
    Document = None

from grounded_daily_monitor import (
    aggregate_records,
    build_open_code_records,
    build_selective_summary,
    extract_future_research_items,
    extract_hypotheses_propositions,
    extract_variable_roles,
)


DESKTOP_ROOT = Path.home() / "Desktop"

PAPER_CODING_SUFFIXES = {".pdf", ".docx", ".txt", ".md"}
META_ANALYSIS_SUFFIXES = {".pdf", ".docx", ".txt", ".md", ".csv", ".xlsx", ".xls"}
INTERVIEW_SUFFIXES = {".docx", ".txt", ".md", ".pdf"}

METHOD_KEYWORDS = [
    "结构方程模型", "sem", "问卷", "survey", "回归", "regression", "元分析", "meta-analysis",
    "案例研究", "case study", "fsqca", "扎根理论", "grounded theory", "访谈", "interview",
    "文本分析", "text analysis", "panel", "面板", "实验", "experiment",
]

THEORY_KEYWORDS = [
    "toe", "技术-组织-环境", "技术—组织—环境", "doi", "创新扩散", "动态能力", "dynamic capability",
    "制度理论", "institutional", "资源基础", "resource-based", "rbv", "社会资本", "social capital",
    "upper echelons", "高层梯队", "utaut", "tam", "可供性", "affordance",
]

SAMPLE_MARKERS = [
    "样本", "问卷", "受访", "respondent", "sample", "manager", "firm", "企业", "公司", "N=",
]

DEFAULT_VARIABLE_PROMPT_TEMPLATE = """我将按照学术文献解构框架，系统分析这篇关于企业人工智能采纳的实证研究论文。

---

## 文献解析报告

### 一、文献标识
请提取：作者、标题、期刊、年份/卷期、在线发表时间、样本特征、分析方法。

### 二、文献类型标注
请判断该文是实证研究、命题研究、综述研究还是方法论文，并简述理论基础与方法论特征。

### 三、变量关系矩阵
请系统提取：
- 自变量
- 中介变量
- 调节变量
- 因变量 / 结果变量
- 控制变量
- 各变量英文表述、测量来源与备注

### 四、假设 / 命题详表
请逐条列出 H1、H2、命题一、命题二等内容，包含理论依据、检验结果、效应方向和显著性。

### 五、模型关系图（文字版）
请把变量之间的关系链写成清晰的文字结构。

### 六、未来研究方向编码表
请提炼未来研究方向，并压缩为词语或短语级编码。

### 七、编码批注与待确认事项
请标记关系不显著、方向矛盾、测量局限、理论解释不足等问题。

### 八、核心贡献与理论对话
请总结理论整合、边界条件、绩效证据、矛盾发现和研究贡献。

请严格依据原文信息，不要编造。

文献信息提示：
- 标题：{title}
- 作者：{authors}
"""


@dataclass
class SourceFile:
    path: str
    name: str
    suffix: str
    size_bytes: int
    estimated_pages: int
    source_root: str
    source_kind: str


@dataclass
class BatchInfo:
    batch_id: str
    file_count: int
    total_size_mb: float
    total_estimated_pages: int
    file_names: list[str]
    file_paths: list[str]


def normalize_space(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("\x00", " ")).strip()


def guess_title(text: str, fallback: str) -> str:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    for line in lines[:30]:
        lowered = line.lower()
        if len(line) < 8 or len(line) > 220:
            continue
        if any(token in lowered for token in ["abstract", "摘要", "关键词", "doi", "vol.", "issue", "journal"]):
            continue
        if re.search(r"(研究|影响|采纳|采用|应用|机制|模型|路径|innovation|adoption|artificial intelligence|ai )", lowered):
            return line
    return fallback


def guess_authors(text: str) -> str:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    for line in lines[:20]:
        if len(line) > 120:
            continue
        if any(token in line.lower() for token in ["university", "学院", "大学", "abstract", "关键词"]):
            continue
        if re.search(r"[A-Z][a-z]+", line) and ("," in line or " and " in line.lower()):
            return line
        if re.search(r"[\u4e00-\u9fff]{2,}", line) and "，" in line:
            return line
    return ""


def guess_year(text: str) -> str:
    match = re.search(r"(20\d{2})", text)
    return match.group(1) if match else ""


def guess_journal(text: str) -> str:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    for line in lines[:25]:
        lowered = line.lower()
        if any(token in lowered for token in ["journal", "review", "management", "research", "forecasting", "technovation", "information"]):
            if len(line) <= 120:
                return line
    return ""


def guess_method(text: str) -> str:
    lowered = text.lower()
    hits = [keyword for keyword in METHOD_KEYWORDS if keyword.lower() in lowered]
    return "、".join(dict.fromkeys(hits))[:200]


def guess_theory(text: str) -> str:
    lowered = text.lower()
    hits = [keyword for keyword in THEORY_KEYWORDS if keyword.lower() in lowered]
    return "、".join(dict.fromkeys(hits))[:200]


def guess_sample(text: str) -> str:
    sentences = [normalize_space(line) for line in text.splitlines() if line.strip()]
    hits: list[str] = []
    for sentence in sentences[:120]:
        lowered = sentence.lower()
        if any(marker.lower() in lowered for marker in SAMPLE_MARKERS) and 12 <= len(sentence) <= 220:
            hits.append(sentence)
        if len(hits) >= 3:
            break
    return "；".join(hits[:2])


def build_prompt(title: str, authors: str, template: str | None = None) -> str:
    base = template or DEFAULT_VARIABLE_PROMPT_TEMPLATE
    return base.format(title=title or "待补充", authors=authors or "待补充")


def ensure_dir(path: Path) -> Path:
    path.mkdir(parents=True, exist_ok=True)
    return path


def list_desktop_directories(limit: int = 80) -> list[Path]:
    if not DESKTOP_ROOT.exists():
        return []
    candidates = [path for path in DESKTOP_ROOT.iterdir() if path.is_dir() and not path.name.startswith(".")]
    return sorted(candidates, key=lambda item: item.name.lower())[:limit]


def normalize_input_paths(raw_paths: list[str]) -> list[Path]:
    paths: list[Path] = []
    seen: set[str] = set()
    for raw in raw_paths:
        if not raw.strip():
            continue
        candidate = Path(raw.strip()).expanduser()
        if not candidate.exists():
            continue
        key = str(candidate.resolve())
        if key in seen:
            continue
        seen.add(key)
        paths.append(candidate.resolve())
    return paths


def estimate_pages(path: Path) -> int:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        try:
            with fitz.open(path) as doc:
                return max(doc.page_count, 1)
        except Exception:
            return max(math.ceil(path.stat().st_size / 180_000), 1)
    if suffix == ".docx" and Document is not None:
        try:
            doc = Document(path)
            chars = sum(len(paragraph.text) for paragraph in doc.paragraphs)
            return max(math.ceil(chars / 1800), 1)
        except Exception:
            pass
    try:
        chars = len(path.read_text(encoding="utf-8", errors="ignore"))
    except Exception:
        chars = max(int(path.stat().st_size / 2), 0)
    return max(math.ceil(chars / 1800), 1)


def scan_source_files(
    input_paths: list[Path],
    *,
    allowed_suffixes: set[str],
    recursive: bool = True,
) -> list[SourceFile]:
    files: list[SourceFile] = []
    seen: set[str] = set()
    for source_path in input_paths:
        candidates = list(source_path.rglob("*")) if source_path.is_dir() and recursive else [source_path]
        for candidate in candidates:
            if not candidate.is_file():
                continue
            if candidate.name.startswith(".~") or candidate.name.startswith("~$"):
                continue
            suffix = candidate.suffix.lower()
            if suffix not in allowed_suffixes:
                continue
            resolved = str(candidate.resolve())
            if resolved in seen:
                continue
            seen.add(resolved)
            stat = candidate.stat()
            files.append(
                SourceFile(
                    path=resolved,
                    name=candidate.name,
                    suffix=suffix,
                    size_bytes=stat.st_size,
                    estimated_pages=estimate_pages(candidate),
                    source_root=str(source_path),
                    source_kind="directory" if source_path.is_dir() else "file",
                )
            )
    return sorted(files, key=lambda item: (item.suffix, item.name.lower()))


def split_into_batches(
    files: list[SourceFile],
    *,
    max_files_per_batch: int = 25,
    max_pages_per_batch: int = 500,
    max_size_mb_per_batch: int = 120,
) -> list[BatchInfo]:
    batches: list[BatchInfo] = []
    current: list[SourceFile] = []
    current_pages = 0
    current_size = 0
    size_limit_bytes = max_size_mb_per_batch * 1024 * 1024

    def flush() -> None:
        nonlocal current, current_pages, current_size
        if not current:
            return
        index = len(batches) + 1
        batches.append(
            BatchInfo(
                batch_id=f"batch_{index:03d}",
                file_count=len(current),
                total_size_mb=round(current_size / (1024 * 1024), 2),
                total_estimated_pages=current_pages,
                file_names=[item.name for item in current],
                file_paths=[item.path for item in current],
            )
        )
        current = []
        current_pages = 0
        current_size = 0

    for item in sorted(files, key=lambda row: (row.estimated_pages, row.size_bytes), reverse=True):
        exceeds_count = current and len(current) >= max_files_per_batch
        exceeds_pages = current and current_pages + item.estimated_pages > max_pages_per_batch
        exceeds_size = current and current_size + item.size_bytes > size_limit_bytes
        if exceeds_count or exceeds_pages or exceeds_size:
            flush()
        current.append(item)
        current_pages += item.estimated_pages
        current_size += item.size_bytes
    flush()
    return batches


def write_inventory(run_dir: Path, files: list[SourceFile], batches: list[BatchInfo]) -> dict[str, str]:
    inventory_dir = ensure_dir(run_dir / "batch_inventory")
    inventory_csv = inventory_dir / "source_inventory.csv"
    batch_csv = inventory_dir / "batch_summary.csv"
    batch_json = inventory_dir / "batch_manifest.json"

    file_to_batch: dict[str, str] = {}
    for batch in batches:
        for path in batch.file_paths:
            file_to_batch[path] = batch.batch_id

    file_rows = []
    for item in files:
        row = asdict(item)
        row["batch_id"] = file_to_batch.get(item.path, "")
        row["size_mb"] = round(item.size_bytes / (1024 * 1024), 2)
        file_rows.append(row)
    pd.DataFrame(file_rows).to_csv(inventory_csv, index=False)
    pd.DataFrame([asdict(batch) for batch in batches]).to_csv(batch_csv, index=False)
    batch_json.write_text(
        json.dumps(
            {
                "total_files": len(files),
                "total_batches": len(batches),
                "batches": [asdict(batch) for batch in batches],
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )
    return {
        "inventory_csv": str(inventory_csv),
        "batch_csv": str(batch_csv),
        "batch_json": str(batch_json),
    }


def build_batch_symlink_folders(run_dir: Path, batches: list[BatchInfo]) -> list[str]:
    batch_root = ensure_dir(run_dir / "batch_inputs")
    batch_dirs: list[str] = []
    for batch in batches:
        batch_dir = ensure_dir(batch_root / batch.batch_id)
        for file_path in batch.file_paths:
            source = Path(file_path)
            target = batch_dir / source.name
            if target.exists():
                continue
            try:
                target.symlink_to(source)
            except Exception:
                shutil.copy2(source, target)
        batch_dirs.append(str(batch_dir))
    return batch_dirs


def extract_text_from_path(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        try:
            with fitz.open(path) as doc:
                return "\n".join(doc.load_page(i).get_text() for i in range(min(doc.page_count, 80)))
        except Exception:
            return ""
    if suffix == ".docx" and Document is not None:
        try:
            doc = Document(path)
            return "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
        except Exception:
            return ""
    try:
        return path.read_text(encoding="utf-8", errors="ignore")
    except Exception:
        return ""


def create_attachment_preview(path: Path, preview_dir: Path) -> str:
    preview_dir.mkdir(parents=True, exist_ok=True)
    preview_path = preview_dir / f"{path.stem[:80]}_{path.suffix.lower().replace('.', '') or 'file'}.png"
    if preview_path.exists():
        return str(preview_path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        try:
            with fitz.open(path) as doc:
                page = doc.load_page(0)
                pix = page.get_pixmap(matrix=fitz.Matrix(0.35, 0.35), alpha=False)
                pix.save(preview_path)
                return str(preview_path)
        except Exception:
            pass
    fig, ax = plt.subplots(figsize=(1.6, 2.1), dpi=160)
    ax.axis("off")
    ax.set_facecolor("#f8fafc")
    fig.patch.set_facecolor("#f8fafc")
    label = f"{path.suffix.lower().replace('.', '').upper() or 'FILE'}\n{path.name[:42]}"
    ax.text(0.5, 0.56, label, ha="center", va="center", fontsize=9, wrap=True, color="#334155")
    ax.text(0.5, 0.13, "附件预览", ha="center", va="center", fontsize=8, color="#0f766e")
    fig.tight_layout(pad=0.2)
    fig.savefig(preview_path, bbox_inches="tight")
    plt.close(fig)
    return str(preview_path)


def split_text_segments(text: str, *, chunk_chars: int = 2200) -> list[str]:
    text = text.replace("\x00", " ").strip()
    if not text:
        return []
    paragraphs = [part.strip() for part in text.splitlines() if part.strip()]
    if not paragraphs:
        return []
    chunks: list[str] = []
    current = ""
    for paragraph in paragraphs:
        candidate = f"{current}\n{paragraph}".strip()
        if current and len(candidate) > chunk_chars:
            chunks.append(current.strip())
            current = paragraph
        else:
            current = candidate
    if current.strip():
        chunks.append(current.strip())
    return chunks


def build_stage1_dataframe(
    run_dir: Path,
    files: list[SourceFile],
    *,
    prompt_template: str | None = None,
) -> pd.DataFrame:
    preview_dir = ensure_dir(run_dir / "stage1_previews")
    rows: list[dict[str, Any]] = []
    for index, item in enumerate(files, start=1):
        path = Path(item.path)
        text = extract_text_from_path(path)[:160_000]
        records = build_open_code_records(text, limit=24)
        matched = aggregate_records(records)
        hypotheses = extract_hypotheses_propositions(text, matched, limit=6)
        variable_roles = extract_variable_roles(text, matched, hypotheses)
        future_items, future_codes = extract_future_research_items(text, matched, limit=6)
        title = guess_title(text, path.stem)
        authors = guess_authors(text)
        year = guess_year(text)
        journal = guess_journal(text)
        method = guess_method(text)
        theory = guess_theory(text)
        sample = guess_sample(text)
        main_concepts = "、".join(
            matched.get("antecedents", [])[:3]
            + matched.get("mechanisms", [])[:2]
            + matched.get("outcomes", [])[:2]
        )
        main_viewpoint = "；".join(hypotheses[:2]) or build_selective_summary(matched)
        rows.append(
            {
                "序号": index,
                "附件预览": create_attachment_preview(path, preview_dir),
                "附件": path.name,
                "文件路径": item.path,
                "批次": "",
                "标题": title,
                "作者": authors or "待补充",
                "期刊": journal or "待补充",
                "年份": year or "待补充",
                "样本特征": sample or "待补充",
                "分析方法": method or "待补充",
                "理论基础": theory or "待补充",
                "主要概念": main_concepts or "待补充",
                "主要观点": main_viewpoint or "待补充",
                "自变量": "、".join(variable_roles.get("independent_vars", [])) or "待补充",
                "中介/调节变量": "、".join(variable_roles.get("mediator_moderator_vars", [])) or "待补充",
                "因变量/结果变量": "、".join(variable_roles.get("dependent_vars", [])) or "待补充",
                "控制变量": "、".join(variable_roles.get("control_vars", [])) or "待补充",
                "未来研究方向": "；".join(future_items) or "待补充",
                "未来研究编码": "、".join(future_codes) or "待补充",
                "变量筛选prompt": build_prompt(title, authors, prompt_template),
            }
        )
    return pd.DataFrame(rows)


def save_stage1_outputs(run_dir: Path, dataframe: pd.DataFrame) -> dict[str, str]:
    output_dir = ensure_dir(run_dir / "stage1_outputs")
    csv_path = output_dir / "paper_stage1_table.csv"
    xlsx_path = output_dir / "paper_stage1_table.xlsx"
    dataframe.to_csv(csv_path, index=False)
    dataframe.to_excel(xlsx_path, index=False)
    return {"csv": str(csv_path), "xlsx": str(xlsx_path)}


def build_meta_analysis_template(run_dir: Path, files: list[SourceFile], batches: list[BatchInfo]) -> str:
    batch_lookup = {}
    for batch in batches:
        for file_path in batch.file_paths:
            batch_lookup[file_path] = batch.batch_id
    rows = []
    for item in files:
        rows.append(
            {
                "batch_id": batch_lookup.get(item.path, ""),
                "filename": item.name,
                "path": item.path,
                "title_guess": Path(item.name).stem,
                "year": "",
                "country": "",
                "sample_size": "",
                "industry": "",
                "independent_var": "",
                "dependent_var": "",
                "mediator": "",
                "moderator": "",
                "effect_size": "",
                "effect_direction": "",
                "method": "",
                "notes": "",
            }
        )
    output = run_dir / "meta_analysis_extraction_template.csv"
    pd.DataFrame(rows).to_csv(output, index=False)
    return str(output)


def build_interview_segments(run_dir: Path, files: list[SourceFile], *, chunk_chars: int = 2200) -> dict[str, str]:
    segment_rows: list[dict[str, Any]] = []
    packet_dir = ensure_dir(run_dir / "coding_packets")
    for item in files:
        path = Path(item.path)
        text = extract_text_from_path(path)
        chunks = split_text_segments(text, chunk_chars=chunk_chars)
        if not chunks:
            continue
        packet_path = packet_dir / f"{path.stem[:80]}.md"
        lines = [f"# {path.name}", ""]
        for index, chunk in enumerate(chunks, start=1):
            segment_id = f"{path.stem}_seg_{index:03d}"
            segment_rows.append(
                {
                    "filename": path.name,
                    "path": str(path),
                    "segment_id": segment_id,
                    "segment_index": index,
                    "text": chunk,
                }
            )
            lines.extend([f"## {segment_id}", "", chunk, ""])
        packet_path.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")
    segment_jsonl = run_dir / "interview_segments.jsonl"
    with segment_jsonl.open("w", encoding="utf-8") as handle:
        for row in segment_rows:
            handle.write(json.dumps(row, ensure_ascii=False) + "\n")
    segment_csv = run_dir / "interview_segments.csv"
    pd.DataFrame(segment_rows).to_csv(segment_csv, index=False)
    return {
        "segment_jsonl": str(segment_jsonl),
        "segment_csv": str(segment_csv),
        "packet_dir": str(packet_dir),
    }
